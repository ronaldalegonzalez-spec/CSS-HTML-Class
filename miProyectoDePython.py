# Import the different utils
import pyinputplus
import re
import time
import openpyxl
import docx

# List of clients
clients = []

# This function displays the main menu of options
def displayMenu():
    print("\nClient Data System")
    print("------------------")
    print("1. Add new client")
    print("2. View clients")
    print("3. Generate report")
    print("4. Search and edit")
    print("5. Exit")
    print("******************")
    return 1


# ===============================================================
"""ADD CLIENTS"""

def addClient():
    name = input("\nEnter name: ")
    name = validateName(name)

    age = pyinputplus.inputInt("Enter age: ", min=18, max=100)

    email = input("Enter email: ")
    while not validateEmail(email):
        print("Invalid email format")
        email = input("Enter email again: ")

    phone = input("Enter phone: ")
    while not validatePhone(phone):
        print("invalid phone number")
        phone = input("Enter phone again: ")
    phone = validatePhone(phone)

    city = input("Enter city: ")
    city = validateCity(city)

    client = name + ", " + str(age) + ", " + email + ", " + phone + ", " + city
    clients.append(client)

    print("\nSaving client...")
    print("Client saved successfully!\n")
    return 1


# ================================================================
"""VALIDATION FUNCTIONS"""

def validateEmail(email):
    pattern = re.compile(r"\S+\@\S+\.\S+")
    patternMatch = re.findall(pattern, email)
    return patternMatch


def validatePhone(phone):
    numberIncorrect = False

    if len(phone) > 10:
        return numberIncorrect
    elif len(phone) < 10:
        return numberIncorrect

    for number in range(len(phone)):
        if not phone[number].isdigit():
            return numberIncorrect

    stylePhone = "(" + phone[:3] + ")" + phone[3:6] + "-" + phone[6:]
    return stylePhone


def validateName(name):
    return name.title()


def validateCity(city):
    return city.title()


# ===============================================================
"""VIEW CLIENTS"""

def viewClients():
    print("\n---------------------------------------")
    print("Client List:")
    print("nº| Name | Age | Email | Phone | City\n")

    if not clients:
        message = '-The list is empty-'
        print(message.rjust(27))
    else:
        for client in range(len(clients)):
            print(str(client+1) + ". " + clients[client])

    print("---------------------------------------")
    return 1


# ===============================================================
"""WRITE REPORT"""

def saveExcel():
    try:
        wb = openpyxl.Workbook()
        sheet = wb.active

        sheet.append(["Name", "Age", "Email", "Phone", "City"])

        for client in clients:
            data = client.split(",")

            for i in range(len(data)):
                data[i] = data[i].strip()

            sheet.append(data)
            
        wb.save("clients.xlsx")
        time.sleep(1)

        print("\nExcel saved successfully!!")
        return 1

    except Exception as e:
        print("\nError saving Excel file!", e)
        return 0


def generateReport():
    totalClients = len(clients)

    cityNames = []
    cityCount = []

    domainNames = []
    domainCount = []

    for client in clients:
        data = client.split(",")

        for i in range(len(data)):
            data[i] = data[i].strip()

        city = data[4]
        email = data[2]

        domain = email[email.index("@")+1:]

        if city not in cityNames:
            cityNames.append(city)
            cityCount.append(1)
        else:
            cityIndex = cityNames.index(city)
            cityCount[cityIndex] += 1

        if domain not in domainNames:
            domainNames.append(domain)
            domainCount.append(1)
        else:
            domainIndex = domainNames.index(domain)
            domainCount[domainIndex] += 1

    try:
        doc = docx.Document()

        doc.add_heading("\t\t       CLIENTS REPORT", 0)
        doc.add_paragraph("\t\t\t          -Summary of customer data-")

        doc.add_heading("Client list:", 1)
        doc.add_paragraph(f"Total clients: {totalClients}")

        doc.add_heading("Clients by city:", 3)
        for i in range(len(cityNames)):
            doc.add_paragraph(f"{cityNames[i]}: {cityCount[i]}")

        doc.add_heading("Email domains:", 3)
        for i in range(len(domainNames)):
            doc.add_paragraph(f"{domainNames[i]}: {domainCount[i]}")

        doc.save("clientReport.docx")
        time.sleep(1)

        print("Report saved successfully!!")
        return 1

    except Exception as e:
        print("Error saving clientReport.docx file!", e)
        return 0


# ===============================================================
"""SEARCH AND EDIT"""

def searchClient():
    print("\n+========================================+")
    print("\t+======================+")
    print("\t\t+======+")
    searchEmail = input("\nEnter client's email to search: ").strip().lower()

    for i in range(len(clients)):
        data = clients[i].split(",")
        email = data[2].strip().lower()

        if email == searchEmail:
            print("\n\t   -*Client found!*-\n")
            return i

    print("\n\t  -*Client not found!*-")
    print("\n\t\t+======+")
    print("\t+======================+")
    print("+========================================+")
    return -1


def editClient(index):
    data = clients[index].split(",")

    for i in range(len(data)):
        data[i] = data[i].strip()

    print(clients[index])
    print("------------------------------------------")
    print("What do you want to edit?\n")
    print("1. Name")
    print("2. Age")
    print("3. Email")
    print("4. Phone")
    print("5. City")
    print("******************")

    choice = input("Choose an option: ")

    if choice == "1":
        data[0] = validateName(input("New name: "))

    elif choice == "2":
        data[1] = str(pyinputplus.inputInt("New age: ", min=18, max=100))

    elif choice == "3":
        email = input("New email: ")
        while not validateEmail(email):
            email = input("Enter valid email: ")
        data[2] = email

    elif choice == "4":
        phone = input("New phone: ")
        while not validatePhone(phone):
            phone = input("Enter valid phone: ")
        data[3] = validatePhone(phone)

    elif choice == "5":
        data[4] = validateCity(input("New city: "))

    clients[index] = data[0] + ", " + data[1] + ", " + data[2] + ", " + data[3] + ", " + data[4]
    
    print("\n\t     *Saving update...")
    time.sleep(1)
    print("\n     -*Client updated successfully!*-")
    print("\n\t\t+======+")
    print("\t+======================+")
    print("+========================================+")

    return 1


# ===============================================================
"""MAIN"""

def main():
    choice = ""

    while choice != "5":
        displayMenu()
        choice = input("Choose an option: ")

        if choice == "1":
            addClient()

        elif choice == "2":
            viewClients()

        elif choice == "3":
            if clients:
                print("Generating excel and report...")
                saveExcel()
                generateReport()
            else:
                print("\n-There's no data to write-")

        elif choice == "4":
            index = searchClient()
            if index != -1:
                editClient(index)

        elif choice == "5":
            print("\nThank you, bye bye!")
            time.sleep(1)

        else:
            print("\nInvalid option!")
    return 1

#call the main fuction
main()
